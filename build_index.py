import os
import pandas as pd
import re
import faiss
import pickle
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from langchain_core.documents import Document as LangchainDocument

# Khởi tạo model embedding
embed_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

# Đường dẫn thư mục chứa file
data_dir = r"C:\Users\Admin\Documents\Python\LLM\Báo cáo tài chính"

# Tên file index và chunks
index_path = "vector.index"
chunks_path = "chunks.pkl"
processed_files_path = "processed_files.pkl"

# Tải FAISS index và các chunks cũ nếu có
if os.path.exists(index_path) and os.path.exists(chunks_path):
    print("✅ Đã có FAISS index. Đang tải...")
    index = faiss.read_index(index_path)
    with open(chunks_path, "rb") as f:
        chunks = pickle.load(f)
    if os.path.exists(processed_files_path):
        with open(processed_files_path, "rb") as f:
            processed_files = pickle.load(f)
    else:
        processed_files = set()
else:
    print("🆕 Chưa có FAISS index. Tạo mới.")
    index = None
    chunks = []
    processed_files = set()

# Hàm đọc file docx
def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join(para.text for para in doc.paragraphs)

# Tổng hợp các văn bản mới
new_paragraphs = []
new_files = []

for file_name in os.listdir(data_dir):
    if file_name in processed_files:
        continue  # Bỏ qua file đã xử lý

    file_path = os.path.join(data_dir, file_name)
    print(f"📄 Đang xử lý: {file_name}")
    try:
        if file_name.endswith(".docx"):
            text = read_docx(file_path)

        elif file_name.endswith(".csv"):
            df = pd.read_csv(file_path)
            text = "\n".join(
                " | ".join([f"{col}: {row[col]}" for col in df.columns if pd.notnull(row[col])])
                for _, row in df.iterrows()
            )

        elif file_name.endswith((".xlsx", ".xls")):
            df = pd.read_excel(file_path)
            text = "\n".join(
                " | ".join([f"{col}: {row[col]}" for col in df.columns if pd.notnull(row[col])])
                for _, row in df.iterrows()
            )

        elif file_name.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        else:
            print(f"⚠️ Bỏ qua file không hỗ trợ: {file_name}")
            continue

        # Làm sạch văn bản
        for line in text.split("\n"):
            cleaned = re.sub(r'[E\[\]L୮]', '', line)
            cleaned = re.sub(r'\s+', ' ', cleaned.strip())
            if cleaned:
                new_paragraphs.append(cleaned)

        new_files.append(file_name)

    except Exception as e:
        print(f"❌ Lỗi xử lý file {file_name}: {e}")

# Nếu có dữ liệu mới thì nhúng và thêm vào FAISS index
if new_paragraphs:
    documents = [LangchainDocument(page_content=p) for p in new_paragraphs]
    splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    new_chunks = splitter.split_documents(documents)
    texts = [chunk.page_content for chunk in new_chunks]
    print(f"🧠 Đang tạo embedding cho {len(texts)} đoạn...")

    embeddings = embed_model.encode(texts, show_progress_bar=True)

    if index is None:
        index = faiss.IndexFlatL2(embeddings.shape[1])
    index.add(embeddings)

    # Cập nhật và lưu lại
    chunks.extend(new_chunks)
    processed_files.update(new_files)

    faiss.write_index(index, index_path)
    with open(chunks_path, "wb") as f:
        pickle.dump(chunks, f)
    with open(processed_files_path, "wb") as f:
        pickle.dump(processed_files, f)

    print(f"✅ Đã thêm {len(new_chunks)} đoạn từ {len(new_files)} file mới.")
else:
    print("🚫 Không có file mới nào cần xử lý.")